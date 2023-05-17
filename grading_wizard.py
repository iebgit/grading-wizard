import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx2txt
import os
import re
import matplotlib.pyplot as plt


# processes answer sheet by creating an ordered list containing only checkboxes and wizards for formatting
def wizard(file):
    doc = docx2txt.process(file)
    checkbox = ['☐', '☒', '🧙']
    checklist ='|'.join(checkbox)
    wiz_list = re.findall(checklist, doc, re.M)
    return wiz_list


# processes answer sheet by creating an ordered list containing only checkboxes for grading
def answers(file):
    doc = docx2txt.process(file)
    checkbox = ['☐', '☒']
    checklist ='|'.join(checkbox)
    ans_list = re.findall(checklist, doc, re.M)
    return ans_list


# processes student paper by creating an ordered list containing only checkboxes
def new(file):
    doc = docx2txt.process(file)
    checkbox = ['☐', '☒']
    checklist = '|'.join(checkbox)
    new_list = re.findall(checklist, doc, re.M)
    return new_list


# compares student paper with answer sheet and generates a graded copy
def writing(file, answer, student, data):

    def add_hyperlink(paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink

    grade = docx.Document(file)
    anList = answer
    stList = student
    cList = []
    iList = []
    count = 0

    for x in data:
        count += 1
        if stList[:x] == anList[:x]:
            cList.append(1)
        else:
            iList.append(count)
        for y in anList[:x]:
            anList.remove(y)
        for z in stList[:x]:
            stList.remove(z)

    numberCorrect = len(cList)
    totalNumber = len(data)

    if len(iList) > 2:
        results = 'The following questions were answered incorrectly: {}\nScore: {}/{}\nPercent: {}%'\
            .format(iList, numberCorrect, totalNumber, (numberCorrect / totalNumber) * 100)
    elif len(iList) == 2:
        results = 'Questions {} and {} were answered incorrectly.\nScore: {}/{}\nPercent: {}%' \
            .format(iList[0], iList[1], numberCorrect, totalNumber, (numberCorrect / totalNumber) * 100)
    elif len(iList) == 1:
        results = 'Question {} was answered incorrectly.\nScore: {}/{}\nPercent: {}%' \
            .format(iList[0], numberCorrect, totalNumber, (numberCorrect / totalNumber) * 100)
    else:
        results = 'No questions were answered incorrectly! 😎\nScore: {}/{}\nPercent: {}%' \
            .format(numberCorrect, totalNumber, (numberCorrect / totalNumber) * 100)

    percent = (numberCorrect / totalNumber) * 100

    if percent > 89:
        grade.add_paragraph('\n🧙 EXCELLENT WORK!')
    elif percent < 60:
        c = 0
        for x in urlList:
            add_hyperlink(grade.add_paragraph("{} --> ".format(labelList[c])), "Follow This Link", x)
            c += 1
    else:
        c = 0
        grade.add_paragraph("\n👌 You're on the right track!")
        for x in urlList:
            add_hyperlink(grade.add_paragraph("{} --> ".format(labelList[c])), "Follow This Link", x)
            c += 1

    grade.add_paragraph(results)

    if percent < 1:
        grade.save("{} [Error].docx".format(file[:-5]))
    else:
        grade.save("{} [Graded].docx".format(file[:-5]))

    return percent


# User interface
c = 0
labelList = []
urlList = []
print('🧙 WELCOME TO THE TEST GRADING WIZARD!')
numURL = int(input("🧙 Number of Practice Links: "))

while c < numURL:
    c += 1
    url = input('🧙 Enter Practice Link: ')
    labelURL = input('🧙 Link Description: ')
    urlList.append(url)
    labelList.append(labelURL)

folderName = input('🧙 Directory: ')
fileName = input("🧙 Answer Sheet: ")
qNum = int(input('🧙 Number of Questions: '))

percentList = []
wizard_list = wizard(fileName)
mcNum = []
i = 0
j = 0

while j < qNum:
    for x in wizard_list:
        if x == '☐' or x == '☒':
            i += 1
            mcNum.insert(j, i)
        elif x == '🧙':
            i = 0
            j += 1

mcNum = mcNum[:qNum]

# iterate through the directory searching for docx files
for folder in os.scandir(folderName):
    if '.docx' in folder.path:
        print('🧙 Success! {}'.format(folder.path))
        perc1 = writing(folder.path, answers(fileName), new(folder.path), mcNum)
        print('{}%'.format(perc1))
        percentList.append(perc1)
    else:
        for file in os.scandir(folder.path):
            if '.docx' in file.path:
                print('🧙 Success! {}'.format(file.path))
                perc2 = writing(file.path, answers(fileName), new(file.path), mcNum)
                print('{}%'.format(perc2))
                percentList.append(perc2)
            else:
                continue

# generate a histogram from student scores
for x in percentList:
    if x == 0:
        percentList.remove(x)

mu = sum(percentList) / len(percentList)
num_bins = 20
n, bins, patches = plt.hist(percentList, num_bins, facecolor='blue', alpha=0.5)
plt.xlabel('Score')
plt.ylabel('Students')
plt.title(r'Histogram of Student Scores: $\mu={}$'.format(mu))
plt.show()
